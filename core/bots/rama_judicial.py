# core/bots/rama_judicial.py
import os
import re
import math
import unicodedata
import asyncio
from datetime import datetime

from django.conf import settings
from asgiref.sync import sync_to_async
from playwright.async_api import async_playwright

import fitz               # PyMuPDF
from docx import Document # python-docx

# ReportLab (render bonito)
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import mm

from core.models import Resultado, Fuente

URL = "https://consultaprocesos.ramajudicial.gov.co/Procesos/NombreRazonSocial"
NOMBRE_SITIO = "rama_judicial"

# ---------------- utils ----------------
def _strip_accents(s: str) -> str:
    s = unicodedata.normalize("NFKD", s or "")
    return "".join(c for c in s if not unicodedata.combining(c))

def _norm(s: str) -> str:
    s = _strip_accents(s)
    s = re.sub(r"[\s\W]+", " ", s or "")
    return s.strip().upper()

def _norm_lower(s: str) -> str:
    s = _strip_accents(s)
    s = re.sub(r"\s+", " ", s or "")
    return s.strip().lower()

def _contains_name(haystack: str, nombre: str) -> bool:
    """match tolerante: todos los tokens del nombre dentro de la línea."""
    H = _norm(haystack)
    tokens = [_norm(t) for t in (nombre or "").split()]
    return all(t in H for t in tokens if t)

def _docx_to_text(path: str) -> str:
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text or "")
        # tablas (si existen)
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join((cell.text or "") for cell in row.cells))
        return "\n".join(parts)
    except Exception:
        return ""

# ---------- NUEVO: helpers para convertir/UNIR ----------
def _pdf_to_png_list(pdf_path: str, out_dir: str, base: str, zoom: float = 2.8) -> list[str]:
    """Convierte cada página del PDF a PNG y devuelve la lista de rutas."""
    os.makedirs(out_dir, exist_ok=True)
    pngs = []
    try:
        with fitz.open(pdf_path) as doc:
            for i in range(doc.page_count):
                pix = doc[i].get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                p = os.path.join(out_dir, f"{base}_p{i+1}.png")
                pix.save(p)
                pngs.append(p)
    except Exception:
        pass
    return pngs

def _merge_pngs_vertical(png_paths: list[str], out_path: str, pad: int = 8, bg=(255, 255, 255)) -> str | None:
    """(Sigue disponible) Une imágenes verticalmente. No se usa por defecto."""
    try:
        from PIL import Image
    except Exception:
        return None

    imgs = [Image.open(p).convert("RGB") for p in png_paths if os.path.exists(p)]
    if not imgs:
        return None

    max_w = max(im.width for im in imgs)
    norm = []
    for im in imgs:
        if im.width != max_w:
            h = int(im.height * max_w / im.width)
            norm.append(im.resize((max_w, h), Image.LANCZOS))
        else:
            norm.append(im)

    total_h = sum(im.height for im in norm) + pad * (len(norm) - 1)
    canvas = Image.new("RGB", (max_w, total_h), bg)
    y = 0
    for im in norm:
        canvas.paste(im, (0, y))
        y += im.height + pad

    canvas.save(out_path, "PNG")
    for im in imgs:
        try: im.close()
        except: pass
    return out_path

def _merge_pngs_grid(png_paths: list[str], out_path: str, cols: int = 2, pad: int = 8, bg=(255, 255, 255)) -> str | None:
    """
    Une varias imágenes en una grilla (cols por fila), manteniendo proporcionalidad.
    Por defecto 2 columnas (dos arriba, dos abajo, ...).
    """
    try:
        from PIL import Image
    except Exception:
        return None

    imgs = [Image.open(p).convert("RGB") for p in png_paths if os.path.exists(p)]
    if not imgs:
        return None

    # Normalizamos todas al mismo ancho destino (el mayor ancho) para alinear columnas
    target_w = max(im.width for im in imgs)
    norm = []
    for im in imgs:
        if im.width != target_w:
            h = int(im.height * target_w / im.width)
            norm.append(im.resize((target_w, h), Image.LANCZOS))
        else:
            norm.append(im)

    rows = math.ceil(len(norm) / max(1, cols))
    # Altura de cada fila = max alto de las imágenes de esa fila
    row_heights = []
    for r in range(rows):
        chunk = norm[r*cols:(r+1)*cols]
        row_heights.append(max(im.height for im in chunk))

    total_w = cols * target_w + pad * (cols - 1)
    total_h = sum(row_heights) + pad * (rows - 1)
    canvas = Image.new("RGB", (total_w, total_h), bg)

    y = 0
    idx = 0
    for r in range(rows):
        x = 0
        for c in range(cols):
            if idx >= len(norm):
                break
            im = norm[idx]
            canvas.paste(im, (x, y))
            x += target_w + pad
            idx += 1
        y += row_heights[r] + pad

    canvas.save(out_path, "PNG")
    for im in imgs:
        try: im.close()
        except: pass
    return out_path

def _pdf_to_single_png(pdf_path: str, out_dir: str, base: str, zoom: float = 2.8) -> list[str]:
    """
    Convierte un PDF a un único PNG. Si hay múltiples páginas,
    las organiza en una grilla (2 columnas por fila).
    Devuelve [ruta_png_final].
    """
    pages = _pdf_to_png_list(pdf_path, out_dir, base, zoom=zoom)
    if not pages:
        return []
    if len(pages) == 1:
        return pages

    merged = os.path.join(out_dir, f"{base}_grid.png")
    if _merge_pngs_grid(pages, merged, cols=2, pad=8):
        return [merged]
    # Fallback: intentar vertical
    merged_v = os.path.join(out_dir, f"{base}_merged.png")
    if _merge_pngs_vertical(pages, merged_v, pad=8):
        return [merged_v]
    # Último recurso, la primera página
    return [pages[0]]

def _text_to_pngs(text: str, out_dir: str, base: str, max_pages: int = 10) -> list[str]:
    """Fallback: crea un PDF simple con el texto y lo convierte a UN solo PNG (grilla 2xN)."""
    os.makedirs(out_dir, exist_ok=True)
    pdf_path = os.path.join(out_dir, f"{base}.pdf")

    c = canvas.Canvas(pdf_path, pagesize=A4)
    w, h = A4
    margin = 40
    y = h - margin
    c.setFont("Helvetica", 10)
    for line in (text or "").splitlines():
        if y < margin + 12:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = h - margin
        c.drawString(margin, y, line[:1800])
        y -= 12
    c.save()

    return _pdf_to_single_png(pdf_path, out_dir, base, zoom=3.0)

# --------- Render bonito: DOCX -> PDF (tablas) -> PNGs ---------
def _docx_to_pretty_pdf(docx_path: str, out_pdf: str):
    doc = Document(docx_path)
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        name="TitleSmall", parent=styles["Heading2"], alignment=1,
        fontSize=12, spaceAfter=8
    )
    body = ParagraphStyle(name="Body8", parent=styles["BodyText"], fontSize=8, leading=10)
    head = ParagraphStyle(name="Head9Bold", parent=styles["BodyText"], fontSize=9, leading=11, spaceAfter=2)

    story = [Paragraph("CONSULTA DE PROCESOS<br/>NACIONAL UNIFICADA", title), Spacer(1, 4)]

    # Primera tabla “grande” (>=5 columnas)
    chosen = None
    for t in doc.tables:
        if t.rows and len(t.rows[0].cells) >= 5 and len(t.rows) >= 2:
            chosen = t
            break

    if not chosen:
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                story.append(Paragraph(txt, body))
        SimpleDocTemplate(
            out_pdf, pagesize=A4, leftMargin=14*mm, rightMargin=14*mm,
            topMargin=12*mm, bottomMargin=12*mm
        ).build(story)
        return

    headers = ["Número de Radicación", "Fecha de Radicación", "Fecha Última Actuac.", "Despacho", "Sujetos Procesales"]

    data = [[Paragraph(h, head) for h in headers]]
    for r in chosen.rows:
        cells = r.cells
        row = []
        for c in cells[:5]:
            txt = (c.text or "").strip().replace("\n", "<br/>")
            row.append(Paragraph(txt, body))
        if any((cells[i].text or "").strip() for i in range(min(5, len(cells)))):  # evita fila de cabecera duplicada
            data.append(row)

    page_w = A4[0]
    usable = page_w - (14*mm + 14*mm)
    col_widths = [32*mm, 25*mm, 28*mm, 72*mm, usable - (32*mm + 25*mm + 28*mm + 72*mm)]

    tbl = Table(data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1b5e20")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,0), 9),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
        ("GRID",       (0,0), (-1,-1), 0.3, colors.grey),
        ("FONTSIZE",   (0,1), (-1,-1), 8),
        ("LEADING",    (0,1), (-1,-1), 10),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F7F7F7")]),
    ]))
    story.append(tbl)

    SimpleDocTemplate(
        out_pdf, pagesize=A4, leftMargin=14*mm, rightMargin=14*mm,
        topMargin=12*mm, bottomMargin=12*mm
    ).build(story)

def _docx_to_pretty_pngs(docx_path: str, out_dir: str, base: str, zoom: float = 3.0) -> list[str]:
    """Devuelve UN solo PNG unido en grilla (2 columnas)."""
    os.makedirs(out_dir, exist_ok=True)
    pdf_path = os.path.join(out_dir, f"{base}.pdf")
    _docx_to_pretty_pdf(docx_path, pdf_path)
    return _pdf_to_single_png(pdf_path, out_dir, base, zoom=zoom)

# --------- NUEVO: leer tabla visible y render a PDF/PNG ---------
async def _scrape_vdatatable_rows(page) -> list[list[str]]:
    """
    Devuelve filas [radicacion, fecha_rad, fecha_ult, despacho, sujetos]
    tomadas de la tabla visible en la página (v-data-table).
    """
    try:
        await page.wait_for_selector("table tbody tr", timeout=8000)
    except Exception:
        return []

    rows = await page.evaluate("""
    () => {
      const out = [];
      const table = document.querySelector("table");
      if (!table) return out;
      const trs = table.querySelectorAll("tbody tr");
      trs.forEach(tr => {
        const tds = tr.querySelectorAll("td");
        if (tds.length < 5) return;

        const rad = (tds[1]?.innerText || "").trim().replace(/\\s+/g, " ");

        let frec = "", fult = "";
        const fechaTxt = (tds[2]?.innerText || "").trim().split("\\n").map(s => s.trim()).filter(Boolean);
        if (fechaTxt.length) {
          frec = fechaTxt[0] || "";
          fult = fechaTxt[1] || "";
        }

        const desp = (tds[3]?.innerText || "").trim().replace(/\\s+/g, " ");

        // Sujetos: mantener saltos de línea, quitar tags <b>, <br>, <div>, etc.
        let sujetosHtml = (tds[4]?.innerHTML || "");
        sujetosHtml = sujetosHtml
          .replace(/<br\\s*\\/?>/gi, "\\n")
          .replace(/<\\/?b[^>]*>/gi, "")
          .replace(/<\\/?span[^>]*>/gi, "")
          .replace(/<\\/?div[^>]*>/gi, "")
          .replace(/&nbsp;/gi, " ")
          .replace(/&amp;/gi, "&");
        const sujetos = sujetosHtml
          .replace(/<[^>]+>/g, "")
          .replace(/\\s+\\n/g, "\\n")
          .replace(/\\n{2,}/g, "\\n")
          .trim();

        if (rad || desp || sujetos) {
          out.push([rad, frec, fult, desp, sujetos]);
        }
      });
      return out;
    }
    """)
    return rows or []

def _rows_to_pdf(rows: list[list[str]], out_pdf: str) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        name="TitleSmall", parent=styles["Heading2"], alignment=1,
        fontSize=14, spaceAfter=10
    )
    body = ParagraphStyle(name="Body8", parent=styles["BodyText"], fontSize=8, leading=10)
    head = ParagraphStyle(name="Head9Bold", parent=styles["BodyText"], fontSize=9, leading=11, spaceAfter=2)

    story = [Paragraph("CONSULTA DE PROCESOS<br/>NACIONAL UNIFICADA", title), Spacer(1, 6)]
    headers = ["Número de Radicación", "Fecha de Radicación", "Fecha Última Actuac.", "Despacho", "Sujetos Procesales"]
    data = [[Paragraph(h, head) for h in headers]]

    for r in rows:
        data.append([
            Paragraph((r[0] or "").replace("\n", "<br/>"), body),
            Paragraph((r[1] or "").replace("\n", "<br/>"), body),
            Paragraph((r[2] or "").replace("\n", "<br/>"), body),
            Paragraph((r[3] or "").replace("\n", "<br/>"), body),
            Paragraph((r[4] or "").replace("\n", "<br/>"), body),
        ])

    page_w = A4[0]
    usable = page_w - (12*mm + 12*mm)
    col_widths = [34*mm, 26*mm, 28*mm, 68*mm, usable - (34*mm + 26*mm + 28*mm + 68*mm)]

    tbl = Table(data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1b5e20")),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,0), 9),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
        ("GRID",       (0,0), (-1,-1), 0.3, colors.grey),
        ("FONTSIZE",   (0,1), (-1,-1), 8),
        ("LEADING",    (0,1), (-1,-1), 10),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F7F7F7")]),
    ]))

    doc = SimpleDocTemplate(
        out_pdf, pagesize=A4, leftMargin=12*mm, rightMargin=12*mm,
        topMargin=12*mm, bottomMargin=12*mm
    )
    story.append(tbl)
    doc.build(story)

def _rows_to_pngs(rows: list[list[str]], out_dir: str, base: str, zoom: float = 2.8) -> list[str]:
    """Render de filas a PDF y devuelve UN solo PNG unido en grilla (2 columnas)."""
    os.makedirs(out_dir, exist_ok=True)
    pdf_path = os.path.join(out_dir, f"{base}.pdf")
    _rows_to_pdf(rows, pdf_path)
    return _pdf_to_single_png(pdf_path, out_dir, base, zoom=zoom)

# --------- UI helpers ----------
async def _check_alert(page):
    """Devuelve texto de alerta (si la hay)."""
    try:
        loc = page.locator(".v-alert__content").first
        if await loc.count():
            txt = await loc.inner_text(timeout=2000)
            return (txt or "").strip()
    except Exception:
        pass
    return None

async def _click_tipo_persona(page, tipo_persona: str):
    opened = False
    for sel in ["#input-72", "div.v-select__selections", "div.v-select input[role='combobox']"]:
        try:
            await page.locator(sel).first.click(timeout=2000)
            opened = True
            break
        except Exception:
            continue
    if not opened:
        try:
            await page.keyboard.press("ArrowDown")
        except Exception:
            pass

    await asyncio.sleep(0.25)
    texto = "Natural" if (tipo_persona or "").strip().lower().startswith("nat") else "Jurídica"
    for sel in [f"div.v-menu__content .v-list-item:has-text('{texto}')", f"text={texto}"]:
        try:
            await page.locator(sel).first.click(timeout=2000)
            return True
        except Exception:
            continue
    return False

async def _fill_nombre(page, nombre: str):
    for sel in ["#input-78", "input#NombreRazonSocial", "form input[type='text']",
                "input[aria-label*='Nombre']", "input[placeholder*='nombre']"]:
        loc = page.locator(sel).first
        if await loc.count():
            await loc.fill(nombre or "")
            return True
    return False

async def _click_opcion_consulta(page, opcion: str = "recientes"):
    opcion = (opcion or "").strip().lower()
    target_id = "#input-194" if opcion.startswith("rec") else "#input-196"
    try:
        await page.locator(target_id).click(timeout=2500)
        return True
    except Exception:
        try:
            label = "Actuaciones Recientes" if opcion.startswith("rec") else "Todos los Procesos"
            await page.locator(f"label:has-text('{label}')").click(timeout=2500)
            return True
        except Exception:
            return False

async def _screenshot_full(page, out_abs: str) -> None:
    try:
        await page.screenshot(path=out_abs, full_page=True)
    except Exception:
        pass

async def _click_volver(page):
    for sel in ["button:has-text('VOLVER')", "button.v-btn:has(span:has-text('Volver'))"]:
        try:
            await page.locator(sel).first.click(timeout=2500)
            return True
        except Exception:
            continue
    return False

# --------------- BOT ---------------
async def consultar_rama_judicial(consulta_id: int, cedula: str, nombre_o_razon: str, tipo_persona: str):
    """
    - Lee la tabla visible (si existe) e incluye Sujetos Procesales en el PNG.
    - 'no generó resultados' -> screenshot + Resultado(score=0).
    - 'Network Error' -> screenshot + mensaje de fallas en la página.
    - 'documento no está disponible' -> Volver y screenshot del listado, score=0.
    - Si no hay tabla -> descargar DOC (flujo original) y render.
    """
    # rutas
    relative_folder = os.path.join("resultados", str(consulta_id))
    absolute_folder = os.path.join(settings.MEDIA_ROOT, relative_folder)
    os.makedirs(absolute_folder, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{NOMBRE_SITIO}_{cedula}_{ts}"
    alert_png_abs = os.path.join(absolute_folder, f"{base}_alert.png")
    alert_png_rel = os.path.join(relative_folder,  f"{base}_alert.png").replace("\\", "/")
    listado_png_abs = os.path.join(absolute_folder, f"{base}_listado.png")
    listado_png_rel = os.path.join(relative_folder,  f"{base}_listado.png").replace("\\", "/")

    fuente_obj = await sync_to_async(lambda: Fuente.objects.filter(nombre=NOMBRE_SITIO).first())()

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            ctx = await browser.new_context(accept_downloads=True, viewport={"width": 1440, "height": 1000})
            page = await ctx.new_page()

            await page.goto(URL, timeout=60000, wait_until="domcontentloaded")
            try:
                await page.wait_for_load_state("networkidle", timeout=10000)
            except Exception:
                pass

            # consulta completa
            await _click_opcion_consulta(page, "todos")
            await _click_tipo_persona(page, tipo_persona)
            if not await _fill_nombre(page, nombre_o_razon):
                raise RuntimeError("No se pudo localizar el campo de nombre/razón social.")

            # enviar
            clicked = False
            for sel in ["button[aria-label='Consultar por nombre o razón social']",
                        "button:has-text('CONSULTAR')", "button:has-text('Consultar')"]:
                try:
                    await page.locator(sel).first.click(timeout=2500)
                    clicked = True
                    break
                except Exception:
                    continue
            if not clicked:
                await page.keyboard.press("Enter")
            await asyncio.sleep(2.5)

            # --- manejar alertas ---
            alert_text = await _check_alert(page)
            if alert_text:
                alert_norm = _norm_lower(alert_text)
                await _screenshot_full(page, alert_png_abs)

                # 0) Network Error
                if "network error" in alert_norm or "error de red" in alert_norm:
                    if fuente_obj:
                        await sync_to_async(Resultado.objects.create)(
                            consulta_id=consulta_id,
                            fuente=fuente_obj,
                            score=0,
                            estado="Sin validar",
                            mensaje="La página presenta fallas (Network Error). Intente más tarde.",
                            archivo=alert_png_rel
                        )
                    await browser.close()
                    return {"mensaje": "network error"}

                # 1) No generó resultados
                if "no genero resultados" in alert_norm:
                    if fuente_obj:
                        await sync_to_async(Resultado.objects.create)(
                            consulta_id=consulta_id,
                            fuente=fuente_obj,
                            score=0,
                            estado="Validado",
                            mensaje="La consulta no generó resultados, por favor revisar las opciones ingresadas e intentarlo nuevamente.",
                            archivo=alert_png_rel
                        )
                    await browser.close()
                    return {"mensaje": "sin resultados"}

                # 2) Documento no disponible → volver y capturar listado
                if "documento solicitado no esta disponible" in alert_norm:
                    await _click_volver(page)
                    for sel in [".v-data-table", "table", ".v-data-table__wrapper"]:
                        try:
                            await page.wait_for_selector(sel, timeout=8000)
                            break
                        except Exception:
                            continue
                    await asyncio.sleep(0.5)
                    await _screenshot_full(page, listado_png_abs)

                    if fuente_obj:
                        await sync_to_async(Resultado.objects.create)(
                            consulta_id=consulta_id,
                            fuente=fuente_obj,
                            score=0,
                            estado="Validado",
                            mensaje="El documento solicitado no está disponible en este momento. Se adjunta listado.",
                            archivo=listado_png_rel if os.path.exists(listado_png_abs) else alert_png_rel
                        )
                    await browser.close()
                    return {"mensaje": "documento no disponible"}

                # 3) Varios registros → volver y seguimos
                if "varios registros" in alert_norm:
                    await _click_volver(page)
                    await asyncio.sleep(0.4)

            # --- Intento 0: leer la tabla visible (incluye Sujetos Procesales) ---
            rows = await _scrape_vdatatable_rows(page)
            if rows:
                pngs_abs = _rows_to_pngs(rows, absolute_folder, f"{base}_render_tabla")
                page_png_rel = os.path.join(relative_folder, os.path.basename(pngs_abs[0])).replace("\\", "/") if pngs_abs else ""

                sujetos_text = "\n".join((r[4] or "") for r in rows)
                es_demandado  = _contains_name(sujetos_text, nombre_o_razon) and "DEMANDADO"   in _norm(sujetos_text)
                es_demandante = _contains_name(sujetos_text, nombre_o_razon) and "DEMANDANTE" in _norm(sujetos_text)
                if es_demandado:
                    score, msg = 10, "El nombre aparece como DEMANDADO/CAUSANTE en uno o más procesos."
                elif es_demandante:
                    score, msg = 6, "El nombre aparece como DEMANDANTE/ACCIONANTE en uno o más procesos."
                else:
                    score, msg = 1, "Se encontraron procesos, pero el nombre no figura como demandante ni demandado."

                await browser.close()
                if fuente_obj:
                    await sync_to_async(Resultado.objects.create)(
                        consulta_id=consulta_id,
                        fuente=fuente_obj,
                        score=score,
                        estado="Validado",
                        mensaje=msg,
                        archivo=page_png_rel
                    )
                return {"mensaje": "ok(tabla)", "score": score, "archivo": page_png_rel}

            # ----- Descargar DOC (flujo original) -----
            doc_rel, page_png_rel, score, msg = "", "", 0, "Se encontraron resultados"
            try:
                async with page.expect_download(timeout=20000) as dl:
                    await page.locator("button:has-text('Descargar DOC')").first.click()
                download = await dl.value

                suggested = download.suggested_filename or f"{base}.docx"
                ext = os.path.splitext(suggested)[1].lower() or ".docx"
                doc_abs = os.path.join(absolute_folder, f"{base}{ext}")
                doc_rel = os.path.join(relative_folder, f"{base}{ext}").replace("\\", "/")

                try:
                    await download.save_as(doc_abs)
                except Exception:
                    tmp = await download.path()
                    if tmp:
                        os.replace(tmp, doc_abs)

                text = ""
                pngs_abs = []
                if ext == ".docx":
                    text = _docx_to_text(doc_abs) or ""
                    pngs_abs = _docx_to_pretty_pngs(doc_abs, absolute_folder, f"{base}_render")

                if not pngs_abs and text:
                    pngs_abs = _text_to_pngs(text, absolute_folder, f"{base}_render")

                page_png_rel = os.path.join(relative_folder, os.path.basename(pngs_abs[0])).replace("\\", "/") if pngs_abs else ""

                # Scoring por sujeto procesal desde texto
                score = 1
                es_demandado  = any(_contains_name(line, nombre_o_razon) and "DEMANDADO"   in _norm(line)
                                    for line in (text or "").splitlines())
                es_demandante = any(_contains_name(line, nombre_o_razon) and "DEMANDANTE" in _norm(line)
                                    for line in (text or "").splitlines())
                if es_demandado:
                    score = 10
                    msg = "El nombre aparece como DEMANDADO/CAUSANTE en uno o más procesos."
                elif es_demandante:
                    score = 6
                    msg = "El nombre aparece como DEMANDANTE/ACCIONANTE en uno o más procesos."
                else:
                    msg = "Se encontraron procesos, pero el nombre no figura como demandante ni demandado."

            except Exception:
                await _screenshot_full(page, alert_png_abs)
                if fuente_obj:
                    await sync_to_async(Resultado.objects.create)(
                        consulta_id=consulta_id,
                        fuente=fuente_obj,
                        score=0,
                        estado="Sin validar",
                        mensaje="No fue posible descargar el DOC de procesos.",
                        archivo=alert_png_rel if os.path.exists(alert_png_abs) else ""
                    )
                await browser.close()
                return {"mensaje": "fallo descarga DOC"}

            await browser.close()

            # ----- Guardar resultado principal -----
            archivo_rel = (page_png_rel or alert_png_rel or doc_rel)
            if fuente_obj:
                await sync_to_async(Resultado.objects.create)(
                    consulta_id=consulta_id,
                    fuente=fuente_obj,
                    score=score,
                    estado="Validado",
                    mensaje=msg,
                    archivo=archivo_rel
                )

            return {"mensaje": "ok", "score": score, "archivo": archivo_rel, "doc": doc_rel}

    except Exception as e:
        fuente_obj = await sync_to_async(lambda: Fuente.objects.filter(nombre=NOMBRE_SITIO).first())()
        if fuente_obj:
            await sync_to_async(Resultado.objects.create)(
                consulta_id=consulta_id,
                fuente=fuente_obj,
                score=0,
                estado="Sin validar",
                mensaje=str(e),
                archivo=""
            )
        return {"mensaje": str(e)}